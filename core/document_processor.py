"""
Модуль для обработки документов различных форматов.
Поддерживает извлечение текста из PDF, DOCX, TXT.
"""

import logging
import io
import re
from typing import Optional, Tuple, Dict, Any
from telegram import Update, Document
from telegram.ext import ContextTypes

# Библиотеки для работы с документами
try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    logging.warning("PyPDF2 не установлен - поддержка PDF отключена")

try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    logging.warning("python-docx не установлен - поддержка DOCX отключена")

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Класс для обработки документов различных форматов."""
    
    # Максимальный размер файла (10 МБ)
    MAX_FILE_SIZE = 10 * 1024 * 1024
    
    # Поддерживаемые форматы
    SUPPORTED_FORMATS = {
        'txt': 'text/plain',
        'pdf': 'application/pdf',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'doc': 'application/msword'
    }
    
    @classmethod
    async def process_document(
        cls, 
        document: Document,
        context: ContextTypes.DEFAULT_TYPE
    ) -> Tuple[bool, str, Optional[str]]:
        """
        Обрабатывает документ и извлекает текст.
        
        Returns:
            Tuple[success: bool, text: str, error: Optional[str]]
        """
        # Проверка размера файла
        if document.file_size > cls.MAX_FILE_SIZE:
            return False, "", f"Файл слишком большой. Максимальный размер: {cls.MAX_FILE_SIZE // (1024*1024)} МБ"
        
        # Определение формата файла
        file_extension = document.file_name.split('.')[-1].lower() if document.file_name else ''
        mime_type = document.mime_type
        
        # Проверка поддержки формата
        if not cls._is_format_supported(file_extension, mime_type):
            supported = ', '.join(cls.SUPPORTED_FORMATS.keys()).upper()
            return False, "", f"Неподдерживаемый формат файла. Поддерживаются: {supported}"
        
        try:
            # Загрузка файла
            file = await document.get_file()
            file_bytes = await file.download_as_bytearray()
            
            # Извлечение текста в зависимости от формата
            if file_extension == 'txt' or mime_type == 'text/plain':
                text = await cls._extract_text_from_txt(file_bytes)
            elif file_extension == 'pdf' or mime_type == 'application/pdf':
                text = await cls._extract_text_from_pdf(file_bytes)
            elif file_extension in ['docx', 'doc'] or 'word' in (mime_type or ''):
                text = await cls._extract_text_from_docx(file_bytes)
            else:
                return False, "", "Не удалось определить формат файла"
            
            # Очистка и валидация текста
            text = cls._clean_text(text)
            
            if not text or len(text.strip()) < 10:
                return False, "", "Документ пустой или содержит слишком мало текста"
            
            # Проверка на слишком большой текст
            if len(text) > 50000:  # Ограничение в 50k символов
                text = text[:50000] + "\n\n[Текст обрезан из-за превышения лимита]"
            
            return True, text, None
            
        except Exception as e:
            logger.error(f"Ошибка обработки документа: {e}", exc_info=True)
            return False, "", f"Ошибка обработки файла: {str(e)}"
    
    @classmethod
    def _is_format_supported(cls, extension: str, mime_type: Optional[str]) -> bool:
        """Проверяет, поддерживается ли формат файла."""
        # Проверка по расширению
        if extension in cls.SUPPORTED_FORMATS:
            return True
        
        # Проверка по MIME-типу
        if mime_type in cls.SUPPORTED_FORMATS.values():
            return True
        
        # Дополнительные проверки для Word документов
        if mime_type and 'word' in mime_type.lower():
            return True
        
        return False
    
    @classmethod
    async def _extract_text_from_txt(cls, file_bytes: bytes) -> str:
        """Извлекает текст из TXT файла."""
        # Пробуем разные кодировки
        encodings = ['utf-8', 'cp1251', 'cp866', 'latin-1']
        
        for encoding in encodings:
            try:
                return file_bytes.decode(encoding)
            except UnicodeDecodeError:
                continue
        
        # Если ничего не подошло, используем utf-8 с игнорированием ошибок
        return file_bytes.decode('utf-8', errors='ignore')
    
    @classmethod
    async def _extract_text_from_pdf(cls, file_bytes: bytes) -> str:
        """Извлекает текст из PDF файла."""
        if not PDF_SUPPORT:
            raise Exception("Поддержка PDF не установлена. Установите PyPDF2: pip install PyPDF2")
        
        text_parts = []
        
        try:
            # Создаем объект PDF
            pdf_file = io.BytesIO(file_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            # Извлекаем текст из каждой страницы
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            return '\n'.join(text_parts)
            
        except Exception as e:
            logger.error(f"Ошибка извлечения текста из PDF: {e}")
            raise Exception(f"Не удалось прочитать PDF файл: {str(e)}")
    
    @classmethod
    async def _extract_text_from_docx(cls, file_bytes: bytes) -> str:
        """Извлекает текст из DOCX файла."""
        if not DOCX_SUPPORT:
            raise Exception("Поддержка DOCX не установлена. Установите python-docx: pip install python-docx")
        
        text_parts = []
        
        try:
            # Создаем объект документа
            doc_file = io.BytesIO(file_bytes)
            doc = DocxDocument(doc_file)
            
            # Извлекаем текст из параграфов
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return '\n'.join(text_parts)
            
        except Exception as e:
            logger.error(f"Ошибка извлечения текста из DOCX: {e}")
            raise Exception(f"Не удалось прочитать DOCX файл: {str(e)}")
    
    @classmethod
    def _clean_text(cls, text: str) -> str:
        """Очищает и нормализует текст."""
        # Удаляем множественные пробелы внутри строк (но сохраняем переносы строк)
        text = re.sub(r'[^\S\n]+', ' ', text)
        # Удаляем множественные пустые строки (3+ → 2)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Удаляем специальные символы и невидимые символы
        text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]', '', text)
        
        # Нормализуем кавычки
        quote_map = {
            "“": '"',
            "”": '"',
            "„": '"',
            "«": '"',
            "»": '"',
            "‘": "'",
            "’": "'",
            "‚": "'",
        }
        text = text.translate(str.maketrans(quote_map))
        
        return text.strip()


class DocumentHandlerMixin:
    """Миксин для добавления обработки документов в обработчики."""
    
    @staticmethod
    async def handle_document_answer(
        update: Update,
        context: ContextTypes.DEFAULT_TYPE,
        task_name: str = "ответ"
    ) -> Optional[str]:
        """
        Обработчик документов для ответов.
        
        Args:
            update: Update объект
            context: Context объект
            task_name: Название задания для сообщений пользователю
            
        Returns:
            Optional[str]: Извлеченный текст или None при ошибке
        """
        if not update.message or not update.message.document:
            await update.message.reply_text("❌ Документ не найден")
            return None
        
        document = update.message.document
        
        # Отправляем сообщение о начале обработки
        processing_msg = await update.message.reply_text(
            "📄 Обрабатываю документ..."
        )
        
        # Обрабатываем документ
        success, text, error = await DocumentProcessor.process_document(
            document, context
        )
        
        # Удаляем сообщение о обработке
        try:
            await processing_msg.delete()
        except Exception as e:
            logger.error("Failed to delete processing message: %s", e)
        
        if not success:
            await update.message.reply_text(
                f"❌ Ошибка обработки документа:\n{error}\n\n"
                f"Пожалуйста, отправьте {task_name} текстовым сообщением."
            )
            return None
        
        # Показываем пользователю, что мы извлекли
        if len(text) > 500:
            preview = text[:500] + "..."
        else:
            preview = text
        
        await update.message.reply_text(
            f"✅ Документ обработан успешно!\n\n"
            f"📝 <b>Извлеченный текст (предпросмотр):</b>\n"
            f"<code>{preview}</code>\n\n"
            f"🔍 Анализирую {task_name}...",
            parse_mode='HTML'
        )
        
        return text
    
    @staticmethod
    def validate_document_content(
        text: str,
        task_type: str = "plan"
    ) -> Tuple[bool, Optional[str]]:
        """
        Валидация содержимого документа для конкретного типа задания.
        
        Returns:
            Tuple[valid: bool, error_message: Optional[str]]
        """
        if task_type == "plan":
            # Проверки для планов (task24)
            if len(text) < 50:
                return False, "План слишком короткий. Минимум 50 символов."
            
            # Проверка на наличие структуры плана
            if not any(marker in text for marker in ['1.', '1)', 'I.', 'а)', 'a)']):
                return False, "Не обнаружена структура плана. Используйте нумерацию пунктов."
            
        elif task_type == "answer":
            # Проверки для обычных ответов
            if len(text) < 10:
                return False, "Ответ слишком короткий."
            
            if len(text) > 10000:
                return False, "Ответ слишком длинный. Максимум 10000 символов."
        
        elif task_type == "examples":
            # Проверки для примеров (task19, task25)
            if len(text) < 30:
                return False, "Примеры слишком короткие."
        
        return True, None


# Функция для простой интеграции в существующие обработчики
async def extract_text_from_document(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE
) -> Optional[str]:
    """
    Простая функция для извлечения текста из документа.
    Возвращает None при ошибке.
    """
    if not update.message or not update.message.document:
        return None
    
    success, text, error = await DocumentProcessor.process_document(
        update.message.document, 
        context
    )
    
    if success:
        return text
    else:
        await update.message.reply_text(f"❌ {error}")
        return None